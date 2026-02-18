import argparse
import base64
import io
import importlib.util
import json
import os
import re
import time
import zipfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import Any, Dict, List, Optional, Tuple
import xml.etree.ElementTree as ET

try:
    import pymupdf as fitz  # PyMuPDF (preferred)
except ModuleNotFoundError:
    import fitz  # type: ignore
import requests
from PIL import Image
from azure.ai.formrecognizer import DocumentAnalysisClient
from azure.core.credentials import AzureKeyCredential
from azure.core.exceptions import HttpResponseError
from dotenv import load_dotenv

try:
    from annotate_pdf_with_essay_rubric import annotate_pdf_essay_pages  # type: ignore
except Exception:
    try:
        parent_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        ann_path = os.path.join(parent_dir, "annotate_pdf_with_essay_rubric.py")
        if os.path.exists(ann_path):
            spec = importlib.util.spec_from_file_location("annotate_pdf_with_essay_rubric", ann_path)
            mod = importlib.util.module_from_spec(spec) if spec else None
            if spec and spec.loader and mod:
                spec.loader.exec_module(mod)
                annotate_pdf_essay_pages = mod.annotate_pdf_essay_pages  # type: ignore
            else:
                annotate_pdf_essay_pages = None  # type: ignore
        else:
            annotate_pdf_essay_pages = None  # type: ignore
    except Exception:
        annotate_pdf_essay_pages = None  # type: ignore


DEFAULT_PRECIS_CRITERIA: List[Dict[str, Any]] = [
    {"id": "comprehension", "criterion": "Comprehension & Understanding of Passage", "marks_allocated": 3},
    {"id": "clarity_expression", "criterion": "Clarity, Expression & Language", "marks_allocated": 3},
    {"id": "brevity", "criterion": "Brevity & Conciseness", "marks_allocated": 2},
    {"id": "organization", "criterion": "Organization & Coherence", "marks_allocated": 2},
    {"id": "tone_meaning", "criterion": "Original Tone & Meaning", "marks_allocated": 2},
    {"id": "originality", "criterion": "Originality & Paraphrasing", "marks_allocated": 2},
    {"id": "grammar_presentation", "criterion": "Grammar & Presentation", "marks_allocated": 1},
    {"id": "title", "criterion": "Title", "marks_allocated": 5},
]

DEFAULT_MODELS: Dict[str, Dict[str, Any]] = {
    "grading": {"model": "grok-4-1-fast-reasoning", "temperature": 0.10},
    "annotations": {"model": "grok-4-1-fast-reasoning", "temperature": 0.15},
    "json_repair": {"model": "grok-4-1-fast-reasoning", "temperature": 0.00},
}

# Keep report text aligned with annotation-style readable text size.
REPORT_BASE_TEXT_SIZE = 12.0

# Populated by _extract_answer_block_text for answer_extracted.json debug output.
_LAST_ANSWER_BLOCK_DEBUG: Dict[str, Any] = {
    "fragmented_handwriting": False,
    "short_ratio": 0.0,
    "body_y_norm": None,
    "line_gap": None,
    "removed_lines": [],
    "kept_lines_count": 0,
    "total_lines_count": 0,
    "used_word_fallback": False,
    "word_count_words": 0,
    "total_words_count": 0,
    "filtered_words_count": 0,
    "title_source": "lines",
}


def _format_duration(seconds: float) -> str:
    if seconds < 60:
        return f"{seconds:.2f}s"
    m = int(seconds // 60)
    s = seconds - m * 60
    return f"{m}m {s:.2f}s"


def clean_json_from_llm(text: str) -> str:
    text = (text or "").strip()
    if not text:
        return ""
    if text.startswith("```"):
        text = re.sub(r"^```[a-zA-Z0-9]*\n?", "", text)
        text = re.sub(r"\n?```$", "", text)
    return text.strip()


def _extract_json_candidate(text: str) -> str:
    s = clean_json_from_llm(text)
    if not s:
        return s
    if s.startswith("{") and s.endswith("}"):
        return s
    if "{" in s and "}" in s:
        start = s.find("{")
        end = s.rfind("}")
        if end > start:
            return s[start : end + 1]
    return s


def _grok_chat(
    grok_api_key: str,
    messages: List[Dict[str, str]],
    model: str = "grok-4-1-fast-reasoning",
    temperature: float = 0.12,
    max_tokens: Optional[int] = None,
    timeout: int = 180,
    max_retries: int = 8,
) -> Dict[str, Any]:
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {grok_api_key}",
    }
    payload: Dict[str, Any] = {
        "model": model,
        "messages": messages,
        "temperature": temperature,
    }
    if max_tokens is not None:
        payload["max_tokens"] = max_tokens

    last_err: Optional[Exception] = None
    for attempt in range(max_retries + 1):
        try:
            resp = requests.post(
                "https://api.x.ai/v1/chat/completions",
                headers=headers,
                json=payload,
                timeout=(30, timeout),
            )
            if resp.status_code >= 300:
                err = RuntimeError(f"Grok API error {resp.status_code}: {resp.text}")
                if resp.status_code in (429, 500, 502, 503, 504) and attempt < max_retries:
                    last_err = err
                    delay = min(60.0, 2.0 ** attempt)
                    print(f"  Grok {resp.status_code} retry {attempt + 1}/{max_retries + 1} in {delay:.0f}s...")
                    time.sleep(delay)
                    continue
                raise err
            return resp.json()
        except (requests.exceptions.Timeout, requests.exceptions.ConnectionError) as e:
            last_err = e
            if attempt >= max_retries:
                raise
            time.sleep(min(60.0, 2.0 ** attempt))

    raise RuntimeError(f"Grok request failed: {last_err}")


def parse_json_with_repair(
    grok_api_key: str,
    raw_text: str,
    *,
    debug_tag: str,
    max_fix_attempts: int = 2,
    repair_model: str = "grok-4-1-fast-reasoning",
    repair_temperature: float = 0.0,
) -> Dict[str, Any]:
    os.makedirs("debug_llm", exist_ok=True)
    with open(os.path.join("debug_llm", f"{debug_tag}_raw.txt"), "w", encoding="utf-8") as f:
        f.write(raw_text or "")

    candidate = _extract_json_candidate(raw_text)
    try:
        return json.loads(candidate)
    except Exception as e:
        last_err = e

    fix_prompt = (
        "Repair the following malformed JSON. Return valid JSON only. "
        "Do not add explanations or markdown."
    )
    current_text = raw_text
    for i in range(max_fix_attempts):
        data = _grok_chat(
            grok_api_key,
            messages=[
                {"role": "system", "content": "You are a JSON repair engine. Return valid JSON only."},
                {"role": "user", "content": fix_prompt + "\n\n" + (current_text or "")},
            ],
            model=repair_model,
            temperature=repair_temperature,
            max_tokens=2500,
        )
        repaired = data["choices"][0]["message"]["content"]
        repaired_candidate = _extract_json_candidate(repaired)
        with open(os.path.join("debug_llm", f"{debug_tag}_repaired_attempt{i+1}.txt"), "w", encoding="utf-8") as f:
            f.write(repaired or "")
        try:
            return json.loads(repaired_candidate)
        except Exception as e:
            last_err = e
            current_text = repaired

    raise ValueError(f"Could not parse JSON after repair attempts: {last_err}")


def _load_docx_text(path: str) -> str:
    """Load text from docx with a pure-XML fallback (no lxml dependency)."""
    try:
        from docx import Document  # type: ignore

        doc = Document(path)
        parts: List[str] = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)
        for t in doc.tables:
            for row in t.rows:
                vals = [((c.text or "").strip().replace("\n", " ")) for c in row.cells]
                vals = [v for v in vals if v]
                if vals:
                    parts.append(" | ".join(vals))
        if parts:
            return "\n".join(parts)
    except Exception:
        pass

    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    with zipfile.ZipFile(path) as z:
        xml = z.read("word/document.xml")
    root = ET.fromstring(xml)
    body = root.find(W + "body")
    parts = []
    if body is None:
        return ""

    def _text(el: ET.Element) -> str:
        return "".join((t.text or "") for t in el.iter(W + "t")).strip()

    for ch in body:
        if ch.tag == W + "p":
            t = _text(ch)
            if t:
                parts.append(t)
        elif ch.tag == W + "tbl":
            for tr in ch.findall(".//" + W + "tr"):
                row_vals: List[str] = []
                for tc in tr.findall(W + "tc"):
                    cell_ps = tc.findall(".//" + W + "p")
                    cell_text = " ".join([_text(p) for p in cell_ps if _text(p)])
                    if cell_text:
                        row_vals.append(cell_text)
                if row_vals:
                    parts.append(" | ".join(row_vals))
    return "\n".join(parts)


def _slugify(s: str) -> str:
    s = re.sub(r"[^a-z0-9]+", "_", (s or "").lower())
    return s.strip("_")[:50] or "criterion"


def parse_precis_rubric_criteria(docx_path: str) -> List[Dict[str, Any]]:
    """Parse criteria + marks from Precis Rubric.docx table."""
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    def _txt(el: ET.Element) -> str:
        return "".join((t.text or "") for t in el.iter(W + "t")).strip()

    rows: List[List[str]] = []
    try:
        with zipfile.ZipFile(docx_path) as z:
            xml = z.read("word/document.xml")
        root = ET.fromstring(xml)
        body = root.find(W + "body")
        if body is None:
            return DEFAULT_PRECIS_CRITERIA

        for tbl in body.findall(W + "tbl"):
            for tr in tbl.findall(".//" + W + "tr"):
                cells: List[str] = []
                for tc in tr.findall(W + "tc"):
                    cell_parts = []
                    for p in tc.findall(".//" + W + "p"):
                        t = _txt(p)
                        if t:
                            cell_parts.append(t)
                    cells.append(" ".join(cell_parts).strip())
                if any(cells):
                    rows.append(cells)
    except Exception:
        return DEFAULT_PRECIS_CRITERIA

    parsed: List[Dict[str, Any]] = []
    for row in rows:
        if len(row) < 2:
            continue
        row_text = " ".join(row)
        if "criterion" in row_text.lower() and "marks" in row_text.lower():
            continue

        marks_raw = row[-1].strip() if row else ""
        m = re.search(r"(\d+(?:\.\d+)?)", marks_raw)
        if not m:
            continue
        marks_val = float(m.group(1))
        if marks_val <= 0:
            continue
        if marks_val > 5:
            if "subtotal" in row_text.lower() or "total" in row_text.lower():
                continue

        crit_raw = row[0].strip()
        if not crit_raw:
            continue
        crit = re.sub(r"^[IVXLCM]+\.?\s*", "", crit_raw, flags=re.IGNORECASE)
        crit = re.sub(r"^\d+\.?\s*", "", crit).strip()
        if not crit:
            continue

        parsed.append(
            {
                "id": _slugify(crit),
                "criterion": crit,
                "marks_allocated": int(round(marks_val)) if abs(marks_val - round(marks_val)) < 1e-6 else marks_val,
            }
        )

    if len(parsed) < 6:
        return DEFAULT_PRECIS_CRITERIA

    parsed = parsed[:8]
    for c in parsed:
        if c["criterion"].lower().startswith("title"):
            c["criterion"] = "Title"
            c["id"] = "title"
    return parsed


def load_environment(env_file: str) -> Tuple[str, DocumentAnalysisClient]:
    load_dotenv(env_file)

    grok_key = os.getenv("Grok_API")
    azure_endpoint = os.getenv("AZURE_ENDPOINT")
    azure_key = os.getenv("AZURE_KEY")

    missing = []
    if not grok_key:
        missing.append("Grok_API")
    if not azure_endpoint:
        missing.append("AZURE_ENDPOINT")
    if not azure_key:
        missing.append("AZURE_KEY")

    if missing:
        raise EnvironmentError(
            f"Missing env vars in {env_file}: {', '.join(missing)}"
        )

    doc_client = DocumentAnalysisClient(
        endpoint=azure_endpoint,
        credential=AzureKeyCredential(azure_key),
    )
    return grok_key, doc_client


def validate_input_paths(pdf_path: str, output_json_path: str, output_pdf_path: str) -> None:
    if not os.path.isfile(pdf_path):
        raise FileNotFoundError(f"PDF not found: {pdf_path}")

    with open(pdf_path, "rb") as f:
        if f.read(4) != b"%PDF":
            raise ValueError(f"Not a valid PDF: {pdf_path}")

    for outp in [output_json_path, output_pdf_path]:
        out_dir = os.path.dirname(outp)
        if out_dir:
            os.makedirs(out_dir, exist_ok=True)


def run_ocr_on_pdf(
    doc_client: DocumentAnalysisClient,
    pdf_path: str,
    *,
    workers: int = 2,
    render_dpi: int = 220,
) -> Dict[str, Any]:
    """Page-wise Azure OCR with retries on payload size."""

    def _encode_page(pil_img: Image.Image, scale: float, quality: int) -> bytes:
        img = pil_img.copy()
        if scale != 1.0:
            img = img.resize((max(1, int(img.width * scale)), max(1, int(img.height * scale))), Image.LANCZOS)
        if img.mode != "RGB":
            img = img.convert("RGB")
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=quality, optimize=True)
        return buf.getvalue()

    def _analyze(img_bytes: bytes) -> Any:
        poller = doc_client.begin_analyze_document("prebuilt-read", document=img_bytes)
        return poller.result()

    doc = fitz.open(pdf_path)
    try:
        pil_pages: List[Tuple[int, Image.Image]] = []
        for idx in range(doc.page_count):
            pix = doc[idx].get_pixmap(dpi=render_dpi)
            pil_pages.append((idx + 1, Image.open(io.BytesIO(pix.tobytes("png")))))
    finally:
        doc.close()

    def _process(page_no: int, pil_img: Image.Image) -> Dict[str, Any]:
        attempts = [(1.0, 75), (0.85, 70), (0.7, 60)]
        result = None
        last_err: Optional[Exception] = None

        for scale, quality in attempts:
            try:
                img_bytes = _encode_page(pil_img, scale=scale, quality=quality)
                result = _analyze(img_bytes)
                break
            except HttpResponseError as e:
                last_err = e
                if "InvalidContentLength" in str(e):
                    continue
                raise

        if result is None:
            raise RuntimeError(f"OCR failed for page {page_no}: {last_err}")

        page_text_parts: List[str] = []
        lines_out: List[Dict[str, Any]] = []
        words_out: List[Dict[str, Any]] = []

        for p in result.pages or []:
            for w in p.words or []:
                wtxt = (w.content or "").strip()
                if not wtxt:
                    continue
                poly = []
                if w.polygon:
                    poly = [(int(pt.x), int(pt.y)) for pt in w.polygon]
                words_out.append(
                    {
                        "text": wtxt,
                        "bbox": poly,
                        "confidence": float(getattr(w, "confidence", 1.0) or 1.0),
                    }
                )
            for ln in p.lines or []:
                ltxt = (ln.content or "").strip()
                if not ltxt:
                    continue
                lpoly = []
                if ln.polygon:
                    lpoly = [(int(pt.x), int(pt.y)) for pt in ln.polygon]
                lines_out.append({"text": ltxt, "bbox": lpoly})
                page_text_parts.append(ltxt)

        return {
            "page_number": page_no,
            "ocr_page_text": " ".join(page_text_parts).strip(),
            "lines": lines_out,
            "words": words_out,
        }

    pages: List[Dict[str, Any]] = []
    worker_count = max(1, int(workers or 1))

    with ThreadPoolExecutor(max_workers=worker_count) as ex:
        futures = {ex.submit(_process, pno, img): pno for pno, img in pil_pages}
        for fut in as_completed(futures):
            pages.append(fut.result())

    pages.sort(key=lambda x: x.get("page_number", 0))
    full_text = "\n".join([p.get("ocr_page_text", "") for p in pages if p.get("ocr_page_text")]).strip()
    return {"pages": pages, "full_text": full_text}


def _is_noise_text(text: str, bbox: List[Tuple[int, int]], page_w: float, page_h: float) -> bool:
    t = (text or "").strip()
    if not t:
        return True
    if len(t) <= 1:
        return True
    if len(re.sub(r"[^A-Za-z0-9]", "", t)) <= 1:
        return True
    if bbox:
        xs = [p[0] for p in bbox]
        ys = [p[1] for p in bbox]
        if xs and ys and page_w > 0 and page_h > 0:
            w = max(xs) - min(xs)
            h = max(ys) - min(ys)
            rel_w = w / max(1.0, page_w)
            rel_h = h / max(1.0, page_h)
            cx = (min(xs) + max(xs)) / 2.0
            cy = (min(ys) + max(ys)) / 2.0
            edge = (cx < page_w * 0.06 or cx > page_w * 0.94 or cy < page_h * 0.06 or cy > page_h * 0.94)
            if rel_w < 0.002 or rel_h < 0.002:
                return True
            if edge and len(t) < 20:
                return True
    return False


def split_extra_text(ocr_data: Dict[str, Any], pdf_path: str) -> Tuple[Dict[str, Any], Dict[str, Any]]:
    """Remove likely watermark/camera/noise OCR lines and return cleaned OCR + extra text pack."""
    page_dims: Dict[int, Tuple[float, float]] = {}
    try:
        doc = fitz.open(pdf_path)
        for i, pg in enumerate(doc):
            r = pg.rect
            page_dims[i + 1] = (float(r.width), float(r.height))
        doc.close()
    except Exception:
        pass

    cleaned_pages: List[Dict[str, Any]] = []
    extra_items: List[Dict[str, Any]] = []

    for p in (ocr_data.get("pages") or []):
        pno = int(p.get("page_number") or 0)
        pw, ph = page_dims.get(pno, (0.0, 0.0))
        kept_lines: List[Dict[str, Any]] = []
        removed_lines: List[Dict[str, Any]] = []
        for ln in (p.get("lines") or []):
            ltxt = (ln.get("text") or "").strip()
            lbbox = ln.get("bbox") or []
            if _is_noise_text(ltxt, lbbox, pw, ph):
                removed_lines.append({"text": ltxt, "bbox": lbbox})
            else:
                kept_lines.append(ln)

        kept_text = " ".join([(x.get("text") or "").strip() for x in kept_lines if (x.get("text") or "").strip()]).strip()
        cp = dict(p)
        cp["lines"] = kept_lines
        cp["ocr_page_text"] = kept_text
        cleaned_pages.append(cp)

        if removed_lines:
            extra_items.append({"page_number": pno, "removed_lines": removed_lines})

    cleaned = {
        "pages": cleaned_pages,
        "full_text": "\n".join([p.get("ocr_page_text", "") for p in cleaned_pages if p.get("ocr_page_text")]).strip(),
    }
    extras = {
        "pdf_path": pdf_path,
        "removed_line_count": sum(len(x.get("removed_lines", [])) for x in extra_items),
        "pages": extra_items,
    }
    return cleaned, extras

def pdf_to_page_images_for_grok(
    pdf_path: str,
    max_pages: Optional[int] = None,
    output_dir: str = "grok_images_precis",
    max_dim: int = 850,
    max_total_base64_chars: int = 280_000,
) -> List[Dict[str, Any]]:
    """Render PDF pages to compact base64 JPEG chunks for Grok."""
    os.makedirs(output_dir, exist_ok=True)

    doc = fitz.open(pdf_path)
    try:
        total_pages = doc.page_count if max_pages is None else min(max_pages, doc.page_count)
        pil_pages: List[Image.Image] = []
        for idx in range(total_pages):
            pix = doc[idx].get_pixmap(dpi=200)
            pil_pages.append(Image.open(io.BytesIO(pix.tobytes("png"))))
    finally:
        doc.close()

    dim_candidates = [max_dim, 720, 640, 560, 512, 448, 384]
    quality_candidates = [68, 60, 52, 45]

    def _encode(dim: int, quality: int, write_files: bool) -> Tuple[List[Dict[str, Any]], int]:
        out: List[Dict[str, Any]] = []
        total = 0
        for i, pil_img in enumerate(pil_pages):
            img = pil_img.copy()
            img.thumbnail((dim, dim))
            if img.mode in ("RGBA", "LA", "P"):
                rgb = Image.new("RGB", img.size, (255, 255, 255))
                if img.mode == "P":
                    img = img.convert("RGBA")
                rgb.paste(img, mask=img.split()[-1] if img.mode in ("RGBA", "LA") else None)
                img = rgb
            elif img.mode != "RGB":
                img = img.convert("RGB")

            buf = io.BytesIO()
            img.save(buf, format="JPEG", quality=quality, optimize=True)
            b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
            total += len(b64)

            file_path = None
            if write_files:
                file_path = os.path.join(output_dir, f"page_{i+1:03d}.jpg")
                with open(file_path, "wb") as f:
                    f.write(buf.getvalue())

            out.append({"page": i + 1, "image_base64": b64, "file_path": file_path, "truncated": False})
        return out, total

    best: Optional[Tuple[List[Dict[str, Any]], int, int, int]] = None
    for dim in dim_candidates:
        for quality in quality_candidates:
            pages, total = _encode(dim, quality, write_files=False)
            best = (pages, total, dim, quality)
            if total > max_total_base64_chars:
                continue
            final_pages, final_total = _encode(dim, quality, write_files=True)
            print(
                f"Saved {len(final_pages)} Grok page images to '{output_dir}/' "
                f"(dim={dim}, quality={quality}, chars={final_total})"
            )
            return final_pages

    if best:
        _, _, dim, quality = best
        final_pages, final_total = _encode(dim, quality, write_files=True)
        print(
            f"Saved {len(final_pages)} Grok page images using fallback settings "
            f"(dim={dim}, quality={quality}, chars={final_total})."
        )
        return final_pages

    return []


def pil_images_to_pdf_bytes(pages: List[Image.Image]) -> bytes:
    out = io.BytesIO()
    if not pages:
        return b""
    pages_rgb = [p.convert("RGB") for p in pages]
    pages_rgb[0].save(out, format="PDF", save_all=True, append_images=pages_rgb[1:])
    return out.getvalue()


def merge_report_and_annotated_answer(
    report_pdf_path: str,
    annotated_pages: List[Image.Image],
    output_pdf_path: str,
) -> None:
    out_doc = fitz.open()
    target_w = 595.0
    target_h = 842.0
    if os.path.exists(report_pdf_path):
        rdoc = fitz.open(report_pdf_path)
        if len(rdoc) > 0:
            r0 = rdoc[0].rect
            target_w, target_h = float(r0.width), float(r0.height)
        out_doc.insert_pdf(rdoc)
        rdoc.close()

    # Keep annotated-answer pages on the SAME page size as report pages.
    for img in annotated_pages:
        page = out_doc.new_page(width=target_w, height=target_h)
        buf = io.BytesIO()
        img.convert("RGB").save(buf, format="PNG")
        stream = buf.getvalue()
        # Preserve aspect ratio; top-align to avoid empty space above annotations.
        img_w, img_h = img.size
        if img_w <= 0 or img_h <= 0:
            continue
        scale = min(target_w / img_w, target_h / img_h)
        draw_w = img_w * scale
        draw_h = img_h * scale
        x0 = (target_w - draw_w) / 2.0
        y0 = 0.0
        rect = fitz.Rect(x0, y0, x0 + draw_w, y0 + draw_h)
        page.insert_image(rect, stream=stream)

    out_doc.save(output_pdf_path)
    out_doc.close()


def _normalize_rating(score: float, max_marks: float) -> str:
    ratio = 0.0 if max_marks <= 0 else score / max_marks
    if ratio >= 0.85:
        return "Excellent"
    if ratio >= 0.70:
        return "Good"
    if ratio >= 0.50:
        return "Average"
    return "Weak"


def _infer_length_status(original_words: int, required_words: int, student_words: int) -> str:
    if original_words <= 0 or required_words <= 0:
        return "Unknown"
    lower = int(round(required_words * 0.95))
    upper = int(round(required_words * 1.05))
    if lower <= student_words <= upper:
        return "Within +/-5%"
    if student_words < lower:
        return "Too Short"
    return "Too Long"


def _page_extent_from_lines(lines: List[Dict[str, Any]]) -> Tuple[float, float]:
    max_x = 0.0
    max_y = 0.0
    for ln in lines:
        poly = ln.get("bbox") or ln.get("polygon") or []
        if isinstance(poly, list) and poly and isinstance(poly[0], (list, tuple)) and len(poly[0]) >= 2:
            pts = poly
        elif isinstance(poly, list) and poly and isinstance(poly[0], dict) and "x" in poly[0]:
            pts = [(p.get("x", 0.0), p.get("y", 0.0)) for p in poly]
        else:
            pts = []
        for x, y in pts:
            max_x = max(max_x, float(x))
            max_y = max(max_y, float(y))
    return max_x, max_y


def _line_bbox_stats(line: Dict[str, Any]) -> Optional[Tuple[float, float, float, float]]:
    poly = line.get("bbox") or line.get("polygon") or []
    if isinstance(poly, list) and poly and isinstance(poly[0], (list, tuple)) and len(poly[0]) >= 2:
        pts = poly
    elif isinstance(poly, list) and poly and isinstance(poly[0], dict) and "x" in poly[0]:
        pts = [(p.get("x", 0.0), p.get("y", 0.0)) for p in poly]
    else:
        pts = []
    if not pts:
        return None
    xs = [float(p[0]) for p in pts]
    ys = [float(p[1]) for p in pts]
    return (min(xs), max(xs), min(ys), max(ys))


def _word_bbox_stats(word: Dict[str, Any]) -> Optional[Tuple[float, float, float, float]]:
    poly = word.get("bbox") or word.get("polygon") or []
    if isinstance(poly, list) and poly and isinstance(poly[0], (list, tuple)) and len(poly[0]) >= 2:
        pts = poly
    elif isinstance(poly, list) and poly and isinstance(poly[0], dict) and "x" in poly[0]:
        pts = [(p.get("x", 0.0), p.get("y", 0.0)) for p in poly]
    else:
        pts = []
    if not pts:
        return None
    xs = [float(p[0]) for p in pts]
    ys = [float(p[1]) for p in pts]
    return (min(xs), max(xs), min(ys), max(ys))


def _page_extent_from_words(words: List[Dict[str, Any]]) -> Tuple[float, float]:
    max_x = 0.0
    max_y = 0.0
    for wd in words:
        bbox = _word_bbox_stats(wd)
        if not bbox:
            continue
        _, x_max, _, y_max = bbox
        max_x = max(max_x, float(x_max))
        max_y = max(max_y, float(y_max))
    return max_x, max_y


def _median(values: List[float]) -> Optional[float]:
    if not values:
        return None
    vs = sorted(values)
    n = len(vs)
    mid = n // 2
    if n % 2 == 1:
        return float(vs[mid])
    return float((vs[mid - 1] + vs[mid]) / 2.0)


def _extract_answer_block_text(ocr_data: Dict[str, Any]) -> Tuple[str, str]:
    global _LAST_ANSWER_BLOCK_DEBUG

    # Page-level handwriting/noisy OCR detector knobs.
    HANDWRITING_SHORT_WORD_MAX = 3
    HANDWRITING_SHORT_CHAR_MAX = 18
    HANDWRITING_FRAGMENTED_RATIO_MIN = 0.45
    HANDWRITING_FEW_LINES_MAX = 8
    HANDWRITING_LATE_BODY_Y_NORM_MIN = 0.20
    HANDWRITING_BODYLIKE_FEW_MAX = 2

    # Top-of-page cleanup knobs (kept explicit for safer future tuning).
    CLEANUP_BODY_Y_MIN_FRAC = 0.12
    CLEANUP_DROP_GAP_MULT = 2.8
    CLEANUP_SHORT_WORD_MAX = 2
    CLEANUP_SHORT_CHAR_MAX = 12

    # Word-level fallback knobs for handwritten/noisy OCR pages.
    WORD_FALLBACK_LINE_COUNT_MAX = 10
    WORD_FALLBACK_ANSWER_WORD_MIN = 25
    WORD_TOP_FILTER_FRAC = 0.06
    WORD_MARGIN_LEFT_FRAC = 0.08
    WORD_MARGIN_RIGHT_FRAC = 0.92
    WORD_ROW_Y_TOL_FRAC = 0.012
    WORD_TITLE_CENTER_TOL_FRAC = 0.18
    WORD_TITLE_MIN_WORDS = 2
    WORD_TITLE_MAX_WORDS = 6
    WORD_TITLE_MAX_Y_NORM = 0.22
    WORD_TITLE_STOPWORDS = {
        "a",
        "an",
        "and",
        "as",
        "at",
        "by",
        "for",
        "from",
        "in",
        "is",
        "of",
        "on",
        "or",
        "the",
        "to",
        "with",
    }

    debug_info: Dict[str, Any] = {
        "fragmented_handwriting": False,
        "short_ratio": 0.0,
        "body_y_norm": None,
        "line_gap": None,
        "removed_lines": [],
        "kept_lines_count": 0,
        "total_lines_count": 0,
        "used_word_fallback": False,
        "word_count_words": 0,
        "total_words_count": 0,
        "filtered_words_count": 0,
        "title_source": "lines",
    }

    pages = ocr_data.get("pages") or []
    page2 = None
    for p in pages:
        if int(p.get("page_number") or 0) == 2:
            page2 = p
            break
    if not page2:
        _LAST_ANSWER_BLOCK_DEBUG = debug_info
        return "", ""

    lines = page2.get("lines") or []
    words_raw = page2.get("words") or []
    page_w, page_h = _page_extent_from_lines(lines)
    if page_w <= 0 or page_h <= 0:
        words_w, words_h = _page_extent_from_words(words_raw)
        page_w = max(page_w, words_w)
        page_h = max(page_h, words_h)
    if page_w <= 0:
        page_w = 1.0
    if page_h <= 0:
        page_h = 1.0

    line_items: List[Dict[str, Any]] = []
    for ln in lines:
        text = (ln.get("text") or ln.get("content") or "").strip()
        if not text:
            continue
        if "improvements" in text.lower():
            continue
        bbox = _line_bbox_stats(ln)
        if bbox:
            x_min, x_max, y_min, y_max = bbox
            center_x = (x_min + x_max) / 2.0
        else:
            x_min = x_max = center_x = None
            y_min = y_max = None
        line_items.append(
            {
                "text": text,
                "word_count": len(text.split()),
                "char_len": len(text),
                "x_min": x_min,
                "x_max": x_max,
                "y_min": y_min if y_min is not None else 0.0,
                "y_max": y_max if y_max is not None else 0.0,
                "center_x": center_x,
            }
        )

    line_items.sort(key=lambda x: x["y_min"])
    debug_info["total_lines_count"] = len(line_items)
    debug_info["total_words_count"] = sum(
        1 for wd in words_raw if str((wd.get("content") or wd.get("text") or "")).strip()
    )

    def _is_body_like(li: Dict[str, Any]) -> bool:
        if li["word_count"] >= 4 or li["char_len"] >= 20:
            cx = li.get("center_x")
            if cx is None:
                return True
            return (0.15 * page_w) <= cx <= (0.85 * page_w)
        return False

    short_ratio = 0.0
    fragmented_handwriting = False
    body_y_norm: Optional[float] = None
    line_gap = 0.04 * page_h
    removed_lines: List[str] = []
    title = ""
    answer_block_text = ""
    title_source = "lines"

    if line_items:
        short_lines_count = sum(
            1
            for li in line_items
            if li["word_count"] <= HANDWRITING_SHORT_WORD_MAX or li["char_len"] <= HANDWRITING_SHORT_CHAR_MAX
        )
        short_ratio = short_lines_count / float(len(line_items))

        body_like = [li for li in line_items if _is_body_like(li)]
        body_start_idx = None
        top_margin = 0.06 * page_h
        for i, li in enumerate(line_items):
            if li["y_min"] < top_margin:
                continue
            if _is_body_like(li):
                body_start_idx = i
                break
        if body_start_idx is None and body_like:
            body_start_idx = line_items.index(body_like[0])
        if body_start_idx is None:
            body_start_idx = 0

        body_y = line_items[body_start_idx]["y_min"]
        body_gaps = []
        prev_y = None
        for li in body_like:
            y = li["y_min"]
            if prev_y is not None:
                dy = y - prev_y
                if dy > 0:
                    body_gaps.append(dy)
            prev_y = y
        measured_gap = _median(body_gaps)
        if measured_gap is not None and measured_gap > 0:
            line_gap = measured_gap

        body_y_norm = float(body_y / page_h) if page_h > 0 else None
        fragmented_handwriting = short_ratio >= HANDWRITING_FRAGMENTED_RATIO_MIN
        if (
            len(line_items) <= HANDWRITING_FEW_LINES_MAX
            and body_y_norm is not None
            and body_y_norm > HANDWRITING_LATE_BODY_Y_NORM_MIN
        ):
            fragmented_handwriting = True
        if len(body_like) <= HANDWRITING_BODYLIKE_FEW_MAX:
            fragmented_handwriting = True

        # Handwritten/noisy OCR pages can split one sentence into many short fragments.
        # On those pages this cleanup is risky, so keep it only for structured layouts.
        if (not fragmented_handwriting) and body_y >= CLEANUP_BODY_Y_MIN_FRAC * page_h:
            drop_y = body_y - CLEANUP_DROP_GAP_MULT * line_gap
            filtered: List[Dict[str, Any]] = []
            for li in line_items:
                if li["y_min"] < drop_y and (
                    li["word_count"] <= CLEANUP_SHORT_WORD_MAX and li["char_len"] <= CLEANUP_SHORT_CHAR_MAX
                ):
                    removed_lines.append(li["text"])
                    continue
                filtered.append(li)
            if len(filtered) >= 2:
                line_items = filtered
            else:
                removed_lines = []

        line_items.sort(key=lambda x: x["y_min"])
        body_like = [li for li in line_items if _is_body_like(li)]
        body_start_idx = None
        for i, li in enumerate(line_items):
            if li["y_min"] < top_margin:
                continue
            if _is_body_like(li):
                body_start_idx = i
                break
        if body_start_idx is None and body_like:
            body_start_idx = line_items.index(body_like[0])
        if body_start_idx is None:
            body_start_idx = 0

        body_y = line_items[body_start_idx]["y_min"]
        body_xs = [li["x_min"] for li in body_like if li["x_min"] is not None]
        body_xe = [li["x_max"] for li in body_like if li["x_max"] is not None]
        body_left = _median(body_xs) if body_xs else None
        body_right = _median(body_xe) if body_xe else None
        body_center = (
            (body_left + body_right) / 2.0 if body_left is not None and body_right is not None else page_w / 2.0
        )
        body_width = (body_right - body_left) if body_left is not None and body_right is not None else page_w
        body_width = max(1.0, body_width)

        title_min_y = body_y - 2.2 * line_gap
        title_max_y = body_y - 0.2 * line_gap
        title_candidates = []
        for li in line_items:
            if li["y_min"] >= body_y:
                continue
            if not (title_min_y <= li["y_min"] <= title_max_y):
                continue
            if li["word_count"] > 7 and li["char_len"] > 30:
                continue
            cx = li.get("center_x")
            if cx is not None:
                if abs(cx - body_center) > (0.25 * body_width):
                    continue
            title_candidates.append(li)

        title_idx = None
        if title_candidates:
            title_candidates.sort(key=lambda x: (-x["y_min"], x["word_count"]))
            chosen = title_candidates[0]
            title = chosen["text"]
            title_idx = line_items.index(chosen)
        else:
            if body_start_idx > 0:
                prev = line_items[body_start_idx - 1]
                if prev["word_count"] <= 7 and (body_y - prev["y_min"]) <= (1.6 * line_gap):
                    title = prev["text"]
                    title_idx = body_start_idx - 1

        body_start = body_start_idx
        if title_idx is not None and title_idx + 1 > body_start:
            body_start = title_idx + 1
        body_lines = [li["text"] for li in line_items[body_start:]]
        answer_block_text = " ".join([l.strip() for l in body_lines if l.strip()]).strip()
    else:
        fragmented_handwriting = True

    debug_info["short_ratio"] = short_ratio
    debug_info["fragmented_handwriting"] = fragmented_handwriting
    debug_info["body_y_norm"] = body_y_norm
    debug_info["line_gap"] = float(line_gap)
    debug_info["removed_lines"] = removed_lines
    debug_info["kept_lines_count"] = len(line_items)

    should_use_word_fallback = fragmented_handwriting and (
        len(line_items) <= WORD_FALLBACK_LINE_COUNT_MAX or len(answer_block_text.split()) < WORD_FALLBACK_ANSWER_WORD_MIN
    )
    used_word_fallback = False
    filtered_words_count = 0
    word_count_words = 0

    if should_use_word_fallback and words_raw:
        word_page_w, word_page_h = _page_extent_from_words(words_raw)
        if page_w <= 1.0 and word_page_w > 0:
            page_w = word_page_w
        if page_h <= 1.0 and word_page_h > 0:
            page_h = word_page_h

        word_items: List[Dict[str, Any]] = []
        for wd in words_raw:
            text = str((wd.get("content") or wd.get("text") or "")).strip()
            if not text:
                continue
            bbox = _word_bbox_stats(wd)
            if not bbox:
                continue
            x_min, x_max, y_min, y_max = bbox
            word_items.append(
                {
                    "text": text,
                    "x_min": x_min,
                    "x_max": x_max,
                    "y_min": y_min,
                    "y_max": y_max,
                    "center_x": (x_min + x_max) / 2.0,
                    "y_center": (y_min + y_max) / 2.0,
                }
            )

        filtered_words: List[Dict[str, Any]] = []
        for wi in word_items:
            if wi["y_min"] < WORD_TOP_FILTER_FRAC * page_h:
                filtered_words_count += 1
                continue
            if wi["center_x"] < WORD_MARGIN_LEFT_FRAC * page_w or wi["center_x"] > WORD_MARGIN_RIGHT_FRAC * page_w:
                filtered_words_count += 1
                continue
            filtered_words.append(wi)

        if filtered_words:
            row_tol = WORD_ROW_Y_TOL_FRAC * page_h
            sorted_words = sorted(filtered_words, key=lambda w: (w["y_center"], w["x_min"]))
            rows: List[Dict[str, Any]] = []
            for wi in sorted_words:
                if rows and abs(wi["y_center"] - rows[-1]["y_center"]) < row_tol:
                    row = rows[-1]
                    row["words"].append(wi)
                    row_n = len(row["words"])
                    row["y_center"] = ((row["y_center"] * (row_n - 1)) + wi["y_center"]) / row_n
                else:
                    rows.append({"y_center": wi["y_center"], "words": [wi]})

            row_items: List[Dict[str, Any]] = []
            for row in rows:
                words_in_row = sorted(row["words"], key=lambda w: w["x_min"])
                row_text = " ".join(w["text"] for w in words_in_row if w["text"]).strip()
                if not row_text:
                    continue
                row_items.append(
                    {
                        "text": row_text,
                        "y_center": row["y_center"],
                        "center_x": sum(w["center_x"] for w in words_in_row) / max(1, len(words_in_row)),
                        "word_count": len(words_in_row),
                    }
                )
            row_items.sort(key=lambda r: r["y_center"])
            word_count_words = sum(int(r["word_count"]) for r in row_items)

            if row_items:
                title_idx = None
                for idx, row in enumerate(row_items):
                    if row["y_center"] / max(1.0, page_h) >= WORD_TITLE_MAX_Y_NORM:
                        continue
                    if abs(row["center_x"] - (page_w / 2.0)) >= WORD_TITLE_CENTER_TOL_FRAC * page_w:
                        continue
                    if not (WORD_TITLE_MIN_WORDS <= row["word_count"] <= WORD_TITLE_MAX_WORDS):
                        continue
                    toks = re.findall(r"[A-Za-z']+", row["text"].lower())
                    if not toks:
                        continue
                    stopword_hits = sum(1 for t in toks if t in WORD_TITLE_STOPWORDS)
                    if stopword_hits > (len(toks) / 2.0):
                        continue
                    title_idx = idx
                    break
                if title_idx is not None:
                    word_title = row_items[title_idx]["text"].strip()
                    word_body_rows = row_items[title_idx + 1 :]
                else:
                    word_title = ""
                    word_body_rows = row_items
                word_body_text = " ".join(r["text"] for r in word_body_rows if r["text"]).strip()
                if word_body_text:
                    title = word_title
                    answer_block_text = word_body_text
                    used_word_fallback = True
                    title_source = "words"

    debug_info["used_word_fallback"] = used_word_fallback
    debug_info["word_count_words"] = int(word_count_words)
    debug_info["filtered_words_count"] = int(filtered_words_count)
    debug_info["title_source"] = title_source

    _LAST_ANSWER_BLOCK_DEBUG = debug_info
    return title.strip(), answer_block_text.strip()


def _extract_question_blocks(ocr_data: Dict[str, Any]) -> Tuple[str, str]:
    pages = ocr_data.get("pages") or []
    page1 = None
    for p in pages:
        if int(p.get("page_number") or 0) == 1:
            page1 = p
            break
    if not page1:
        return "", ""

    lines = page1.get("lines") or []
    words = page1.get("words") or []

    def _norm_ws(text: str) -> str:
        return re.sub(r"\s+", " ", (text or "")).strip()

    # Anchor-based parsing is needed because noisy scans can lose short instruction lines
    # after line-level cleanup; word rows preserve page-1 prompt text more reliably.
    line_items: List[Tuple[float, str]] = []
    if words:
        words_h = _page_extent_from_words(words)[1]
        lines_h = _page_extent_from_lines(lines)[1]
        page_h = max(words_h, lines_h, 1.0)
        row_tol = 0.012 * page_h
        word_items: List[Dict[str, Any]] = []
        for wd in words:
            text = str((wd.get("text") or wd.get("content") or "")).strip()
            if not text:
                continue
            bbox = _word_bbox_stats(wd)
            if not bbox:
                continue
            x_min, _, y_min, y_max = bbox
            word_items.append({"text": text, "x_min": x_min, "y_center": (y_min + y_max) / 2.0})
        word_items.sort(key=lambda w: (w["y_center"], w["x_min"]))

        rows: List[Dict[str, Any]] = []
        for wi in word_items:
            if rows and abs(wi["y_center"] - rows[-1]["y_center"]) < row_tol:
                row = rows[-1]
                row["words"].append(wi)
                n = len(row["words"])
                row["y_center"] = ((row["y_center"] * (n - 1)) + wi["y_center"]) / n
            else:
                rows.append({"y_center": wi["y_center"], "words": [wi]})

        for row in rows:
            words_in_row = sorted(row["words"], key=lambda w: w["x_min"])
            row_text = _norm_ws(" ".join(w["text"] for w in words_in_row if w["text"]))
            if row_text:
                line_items.append((float(row["y_center"]), row_text))
    else:
        for idx, ln in enumerate(lines):
            text = (ln.get("text") or ln.get("content") or "").strip()
            if not text:
                continue
            bbox = _line_bbox_stats(ln)
            y = float(bbox[2]) if bbox else float(idx)
            line_items.append((y, text))

    line_items.sort(key=lambda x: x[0])
    if not line_items:
        return "", ""

    page1_text = _norm_ws("\n".join([t for _, t in line_items if t]))
    if not page1_text:
        return "", ""

    instruction_anchor_patterns = [
        r"\bexercise\b",
        r"write\s+a\s+pr[eé]cis",
        r"\bpr[eé]cis\b",
        r"suggest\s+a\s+suitable\s+title",
    ]
    passage_start_idx: Optional[int] = None
    m_drop = re.search(r"\ba\s+drop\s+of\s+water\b", page1_text, flags=re.IGNORECASE)
    if m_drop:
        passage_start_idx = m_drop.start()
    else:
        m_alas = re.search(r"[\"“”']?\s*alas!?", page1_text, flags=re.IGNORECASE)
        if m_alas:
            passage_start_idx = m_alas.start()

    if passage_start_idx is not None:
        instr_positions: List[int] = []
        for pat in instruction_anchor_patterns:
            for m in re.finditer(pat, page1_text, flags=re.IGNORECASE):
                if m.start() < passage_start_idx:
                    instr_positions.append(m.start())
        q_start = min(instr_positions) if instr_positions else 0
        question_text = _norm_ws(page1_text[q_start:passage_start_idx]).strip(" :-")
        statement_text = _norm_ws(page1_text[passage_start_idx:])
        if not question_text:
            question_text = _norm_ws(page1_text[:passage_start_idx]).strip(" :-")
        return question_text, statement_text

    # Fallback when anchors are missing.
    question_lines: List[str] = []
    statement_lines: List[str] = []
    in_statement = False
    keywords = ("exercise", "precis", "title", "question", "write")
    for _, text in line_items:
        t_low = text.lower()
        word_count = len(text.split())
        has_keyword = any(k in t_low for k in keywords)
        passage_like = word_count >= 8 and not has_keyword
        if not in_statement and passage_like:
            in_statement = True
            statement_lines.append(text)
            continue
        if in_statement:
            statement_lines.append(text)
        else:
            question_lines.append(text)
    return _norm_ws(" ".join(question_lines)), _norm_ws(" ".join(statement_lines))


def call_grok_for_precis_grading(
    grok_api_key: str,
    rubric_text: str,
    criteria_template: List[Dict[str, Any]],
    ocr_data: Dict[str, Any],
    page_images: List[Dict[str, Any]],
    *,
    model: str = "grok-4-1-fast-reasoning",
    temperature: float = 0.10,
    repair_model: str = "grok-4-1-fast-reasoning",
    repair_temperature: float = 0.0,
) -> Dict[str, Any]:
    def _to_int(value: Any, default: int = 0) -> int:
        try:
            return int(round(float(value)))
        except Exception:
            return int(default)

    total_marks = int(round(sum(float(c.get("marks_allocated", 0)) for c in criteria_template)))
    answer_title_hint, answer_block_text = _extract_answer_block_text(ocr_data)
    answer_debug = _LAST_ANSWER_BLOCK_DEBUG if isinstance(_LAST_ANSWER_BLOCK_DEBUG, dict) else {}
    answer_block_word_count = len((answer_block_text or "").split())
    answer_total_lines_count = int(answer_debug.get("total_lines_count", 0) or 0)
    answer_fragmented_handwriting = bool(answer_debug.get("fragmented_handwriting", False))
    answer_used_word_fallback = bool(answer_debug.get("used_word_fallback", False))
    answer_word_count_words = int(answer_debug.get("word_count_words", 0) or 0)
    question_block_text, question_statement_block_text = _extract_question_blocks(ocr_data)
    answer_page_images = [p for p in page_images if int(p.get("page") or 0) == 2]
    if not answer_page_images:
        answer_page_images = page_images

    schema_hint = {
        "question_text": "",
        "question_statement_text": "",
        "topic": "",
        "student_title": "",
        "student_precis_text": "",
        "original_passage_word_count": 0,
        "required_precis_word_count": 0,
        "student_precis_word_count": 0,
        "length_status": "Within +/-5% | Too Long | Too Short | Unknown",
        "criteria": [
            {
                "id": c.get("id"),
                "criterion": c.get("criterion"),
                "marks_allocated": c.get("marks_allocated"),
                "marks_awarded": 0,
                "rating": "Weak",
                "key_comments": "",
            }
            for c in criteria_template
        ],
        "total_awarded": 0,
        "overall_rating": "Weak",
        "reasons_for_low_score": [""],
        "ideal_precis": {
            "title": "",
            "text": "",
        },
        "overall_remarks": "",
    }

    system = {
        "role": "system",
        "content": (
            "You are a strict CSS precis examiner. "
            "Return valid JSON only with no markdown or commentary."
        ),
    }

    instructions = (
        "Evaluate a 2-page precis submission.\n"
        "Document layout expectation:\n"
        "- Page 1 contains question/prompt and original source passage.\n"
        "- Page 2 contains student answer: title + one precis paragraph.\n"
        "Required length rule:\n"
        "- Required precis length is exactly one-third of original passage word count.\n"
        "- Also classify length_status using +/-5% tolerance around required length.\n"
        "Scoring rules:\n"
        "- Follow the provided rubric criteria and marks exactly.\n"
        "- For each criterion, give marks_awarded as a whole number within [0, marks_allocated].\n"
        "- Add concise, evidence-based key_comments for each criterion.\n"
        "- Provide total_awarded as the sum of marks_awarded values.\n"
        "- Provide overall_rating from: Excellent, Good, Average, Weak.\n"
        "- reasons_for_low_score must contain only concrete weaknesses, not praise.\n"
        "- reasons_for_low_score must NOT contain 'No major weaknesses identified'.\n"
        "- If total_awarded is below 70% of total marks, reasons_for_low_score must contain at least 2 concrete items.\n"
        "- ideal_precis.text must be a high-quality improved precis for this same passage.\n"
        "- ideal_precis.title must be concise and relevant.\n"
        "- ideal_precis.title and ideal_precis.text must both be non-empty.\n"
        "Question extraction rules:\n"
        "- question_text must come from question_block_text only.\n"
        "- question_statement_text must come from question_statement_block_text only.\n"
        "Extraction rules:\n"
        "- topic should be the passage topic/theme.\n"
        "- student_title must be exactly what the student wrote when visible; else infer carefully.\n"
        "- student_precis_text must contain the student precis paragraph only.\n"
        "- Use answer_title_hint and answer_block_text as primary evidence for the student's title/precis.\n"
        "- Those hints are OCR-filtered to remove stray top-of-page text from the previous paragraph.\n"
        "- You MUST transcribe student_precis_text from page 2 image only when extracted text is insufficient.\n"
        "- Insufficient means: answer_block_text has fewer than 20 words OR word_count_words < 25 OR (used_word_fallback is false AND total_lines_count < 10).\n"
        "- If used_word_fallback is true and word_count_words >= 35, DO NOT force page-image transcription; use answer_block_text as authoritative.\n"
        "- Compute original_passage_word_count, required_precis_word_count, student_precis_word_count.\n"
        "- Ignore unrelated watermark/camera/footer artifacts or stray words not part of question/answer content.\n"
        "- Do not mention OCR or handwriting quality in comments.\n"
        "Return JSON strictly matching the schema."
    )

    payload = {
        "rubric_text": rubric_text,
        "criteria_template": criteria_template,
        "question_block_text": question_block_text,
        "question_statement_block_text": question_statement_block_text,
        "answer_title_hint": answer_title_hint,
        "answer_block_text": answer_block_text,
        "answer_block_word_count": answer_block_word_count,
        "total_lines_count": answer_total_lines_count,
        "fragmented_handwriting": answer_fragmented_handwriting,
        "used_word_fallback": answer_used_word_fallback,
        "word_count_words": answer_word_count_words,
        "ocr_pages": ocr_data.get("pages", []),
        "ocr_full_text": ocr_data.get("full_text", ""),
        "answer_page_images": answer_page_images,
        "page_images": page_images,
        "output_schema": schema_hint,
    }

    def _validate(parsed: Dict[str, Any]) -> bool:
        if not isinstance(parsed, dict):
            return False

        crit = parsed.get("criteria")
        crit_list: List[Dict[str, Any]] = crit if isinstance(crit, list) else []
        marks_sum = 0
        normalized_criteria: List[Dict[str, Any]] = []
        for i, tmpl in enumerate(criteria_template):
            c = crit_list[i] if i < len(crit_list) and isinstance(crit_list[i], dict) else {}
            alloc = float(criteria_template[i].get("marks_allocated", 0))
            try:
                aw = float(c.get("marks_awarded", 0))
            except Exception:
                aw = 0.0
            aw = min(max(aw, 0.0), alloc)
            aw_int = int(round(aw))
            rating = str(c.get("rating", "") or "").strip()
            if rating not in ("Excellent", "Good", "Average", "Weak"):
                rating = _normalize_rating(float(aw_int), alloc)
            normalized_criteria.append(
                {
                    "id": str(tmpl.get("id", "") or f"criterion_{i+1}"),
                    "criterion": str(tmpl.get("criterion", "") or ""),
                    "marks_allocated": int(round(alloc)),
                    "marks_awarded": aw_int,
                    "rating": rating,
                    "key_comments": str(c.get("key_comments", "") or "").strip(),
                }
            )
            marks_sum += aw_int
        parsed["criteria"] = normalized_criteria

        declared_total = parsed.get("total_awarded")
        try:
            declared_total_f = float(declared_total)
        except Exception:
            declared_total_f = float(marks_sum)

        if abs(declared_total_f - marks_sum) > 0.75:
            parsed["total_awarded"] = int(round(marks_sum))
        else:
            parsed["total_awarded"] = int(round(declared_total_f))

        parsed["total_awarded"] = int(max(0, min(int(total_marks), int(parsed["total_awarded"]))))

        if parsed.get("overall_rating") not in ("Excellent", "Good", "Average", "Weak"):
            parsed["overall_rating"] = _normalize_rating(float(parsed["total_awarded"]), float(total_marks))

        # Keep question fields anchored to extracted page-1 blocks for consistency.
        parsed["question_text"] = question_block_text
        parsed["question_statement_text"] = question_statement_block_text

        parsed["topic"] = str(parsed.get("topic", "") or "").strip()
        parsed["student_title"] = str(parsed.get("student_title", "") or "").strip()
        parsed["student_precis_text"] = str(parsed.get("student_precis_text", "") or "").strip()
        if not parsed["student_title"]:
            parsed["student_title"] = str(answer_title_hint or "").strip()
        if not parsed["student_precis_text"]:
            parsed["student_precis_text"] = str(answer_block_text or "").strip()

        ow = _to_int(parsed.get("original_passage_word_count"), 0)
        rw = _to_int(parsed.get("required_precis_word_count"), 0)
        sw = len(parsed["student_precis_text"].split())
        parsed["student_precis_word_count"] = sw
        if ow <= 0 and question_statement_block_text:
            ow = len(question_statement_block_text.split())
            parsed["original_passage_word_count"] = ow
        if rw <= 0 and ow > 0:
            parsed["required_precis_word_count"] = int(round(ow / 3.0))
            rw = int(parsed["required_precis_word_count"])

        if not parsed.get("length_status") or str(parsed.get("length_status")).strip() == "Unknown":
            parsed["length_status"] = _infer_length_status(ow, rw, sw)

        reasons = parsed.get("reasons_for_low_score")
        if not isinstance(reasons, list):
            parsed["reasons_for_low_score"] = []
        parsed["reasons_for_low_score"] = [str(x).strip() for x in parsed["reasons_for_low_score"] if str(x).strip()][:8]

        ideal = parsed.get("ideal_precis")
        if not isinstance(ideal, dict):
            parsed["ideal_precis"] = {"title": "", "text": ""}
        parsed["ideal_precis"]["title"] = str(parsed["ideal_precis"].get("title", "")).strip()
        parsed["ideal_precis"]["text"] = str(parsed["ideal_precis"].get("text", "")).strip()
        parsed["overall_remarks"] = str(parsed.get("overall_remarks", "") or "").strip()

        # Quality gates: reject low-quality but schema-valid outputs and force retry.
        all_marks_zero = all(int(c.get("marks_awarded", 0)) <= 0 for c in parsed["criteria"])
        if all_marks_zero:
            return False
        if not parsed["ideal_precis"]["text"]:
            return False
        reasons_lc = [str(x).strip().lower() for x in parsed["reasons_for_low_score"] if str(x).strip()]
        if not reasons_lc:
            return False
        if any("no major weaknesses" in r for r in reasons_lc):
            return False
        if sw >= 30 and int(parsed.get("total_awarded", 0) or 0) == 0:
            return False
        return True

    last_err: Optional[Exception] = None
    for attempt in range(4):
        print(f"  Precis grading attempt {attempt + 1}/4...")
        response = _grok_chat(
            grok_api_key,
            messages=[
                system,
                {
                    "role": "user",
                    "content": instructions + "\n\nDATA:\n" + json.dumps(payload, ensure_ascii=False),
                },
            ],
            model=model,
            temperature=temperature,
            max_tokens=4000,
        )
        content = response["choices"][0]["message"]["content"]
        try:
            parsed = parse_json_with_repair(
                grok_api_key,
                content,
                debug_tag="precis_grading",
                max_fix_attempts=3,
                repair_model=repair_model,
                repair_temperature=repair_temperature,
            )
        except Exception as e:
            last_err = e
            continue

        if _validate(parsed):
            print(f"  Precis grading validated on attempt {attempt + 1}.")
            return parsed
        last_err = ValueError("Invalid grading JSON")
    print(f"  Warning: grading JSON remained incomplete after retries ({last_err}); using fallback normalization.")
    fallback_student_title = str(answer_title_hint or "").strip() or "Precis"
    fallback_student_text = str(answer_block_text or "").strip()
    fallback_ow = len(question_statement_block_text.split()) if question_statement_block_text else 0
    fallback_rw = int(round(fallback_ow / 3.0)) if fallback_ow > 0 else 0
    fallback_sw = len(fallback_student_text.split()) if fallback_student_text else 0

    source_text = str(question_statement_block_text or fallback_student_text).strip()
    source_sents = [s.strip() for s in re.split(r"(?<=[.!?])\s+", source_text) if s.strip()]
    if not source_sents and source_text:
        source_sents = [source_text]
    if source_sents:
        ideal_text_parts: List[str] = []
        running_words = 0
        max_words = max(55, fallback_rw) if fallback_rw > 0 else 70
        for s in source_sents:
            swc = len(s.split())
            if running_words > 0 and running_words + swc > max_words:
                break
            ideal_text_parts.append(s)
            running_words += swc
            if len(ideal_text_parts) >= 4:
                break
        fallback_ideal_text = " ".join(ideal_text_parts).strip()
    else:
        fallback_ideal_text = ""
    if not fallback_ideal_text:
        fallback_ideal_text = (
            "The passage highlights personal significance, humility, and growth through contribution to others."
        )

    title_seed = str(answer_title_hint or "").strip()
    if title_seed:
        fallback_ideal_title = title_seed
    else:
        title_tokens = re.findall(r"[A-Za-z]+", source_text)
        if len(title_tokens) >= 3:
            fallback_ideal_title = " ".join(title_tokens[:3]).title()
        elif title_tokens:
            fallback_ideal_title = title_tokens[0].title()
        else:
            fallback_ideal_title = "Precis"

    fallback_reasons: List[str] = []
    if fallback_rw > 0 and fallback_sw > 0:
        if fallback_sw < int(round(fallback_rw * 0.95)):
            fallback_reasons.append("Response is shorter than required one-third length and omits key supporting points.")
        elif fallback_sw > int(round(fallback_rw * 1.05)):
            fallback_reasons.append("Response exceeds the required one-third length and needs tighter compression.")
    fallback_reasons.append("Main ideas are only partially developed, so logical progression and clarity need improvement.")
    fallback_reasons.append("Language precision and paraphrasing control need refinement to preserve meaning without copying.")
    fallback_reasons = [r for r in fallback_reasons if r][:3]

    fallback_criteria: List[Dict[str, Any]] = []
    fallback_marks_sum = 0
    for i, tmpl in enumerate(criteria_template):
        alloc = int(round(float(tmpl.get("marks_allocated", 0) or 0)))
        if alloc <= 0:
            aw = 0
        elif fallback_sw > 0:
            aw = max(1, int(round(alloc * 0.4)))
            aw = min(alloc, aw)
        else:
            aw = 0
        key_comment = "Partial achievement; response needs clearer coverage and tighter wording."
        crit_name = str(tmpl.get("criterion", "") or "").lower()
        if "title" in crit_name:
            key_comment = "Title is present but can be more specific and aligned with the central idea."
        elif "brevity" in crit_name:
            key_comment = "Compression is uneven; reduce peripheral detail and keep core points concise."
        elif "organization" in crit_name:
            key_comment = "Organization is basic; transitions and sentence flow need improvement."
        fallback_criteria.append(
            {
                "id": str(tmpl.get("id", "") or f"criterion_{i+1}"),
                "criterion": str(tmpl.get("criterion", "") or ""),
                "marks_allocated": alloc,
                "marks_awarded": aw,
                "rating": _normalize_rating(float(aw), float(alloc) if alloc > 0 else 1.0),
                "key_comments": key_comment,
            }
        )
        fallback_marks_sum += aw

    if fallback_marks_sum == 0 and fallback_criteria:
        first_alloc = int(fallback_criteria[0].get("marks_allocated", 0) or 0)
        if first_alloc > 0:
            fallback_criteria[0]["marks_awarded"] = 1
            fallback_criteria[0]["rating"] = _normalize_rating(1.0, float(first_alloc))
            fallback_marks_sum = 1

    fallback: Dict[str, Any] = {
        "question_text": question_block_text,
        "question_statement_text": question_statement_block_text,
        "topic": "",
        "student_title": fallback_student_title,
        "student_precis_text": fallback_student_text,
        "original_passage_word_count": fallback_ow,
        "required_precis_word_count": fallback_rw,
        "student_precis_word_count": fallback_sw,
        "length_status": _infer_length_status(fallback_ow, fallback_rw, fallback_sw),
        "criteria": fallback_criteria,
        "total_awarded": int(fallback_marks_sum),
        "overall_rating": _normalize_rating(float(fallback_marks_sum), float(total_marks) if total_marks > 0 else 1.0),
        "reasons_for_low_score": fallback_reasons[:],
        "ideal_precis": {"title": fallback_ideal_title, "text": fallback_ideal_text},
        "overall_remarks": "Fallback grading used after repeated invalid model outputs; scores are conservative.",
    }
    if not _validate(fallback):
        fallback["reasons_for_low_score"] = [
            "Coverage is incomplete and several key points from the source are missing.",
            "Expression needs tighter paraphrasing and clearer organization.",
        ]
        fallback["ideal_precis"] = {
            "title": fallback_ideal_title or "Precis",
            "text": fallback_ideal_text or "The passage emphasizes humility, purpose, and value through contribution.",
        }
    return fallback


def call_grok_for_precis_annotations(
    grok_api_key: str,
    annotations_rubric_text: str,
    ocr_data: Dict[str, Any],
    grading: Dict[str, Any],
    page_images: List[Dict[str, Any]],
    *,
    model: str = "grok-4-1-fast-reasoning",
    temperature: float = 0.15,
    repair_model: str = "grok-4-1-fast-reasoning",
    repair_temperature: float = 0.0,
) -> Dict[str, Any]:
    """
    Returns:
    {
      "page_suggestions": [{"page": 2, "suggestions": ["..."]}],
      "annotations": [ ... ],
      "errors": [ ... ]
    }
    """
    system = {
        "role": "system",
        "content": (
            "You generate precise, locatable annotations for handwritten precis scripts.\n"
            "Primary truth is page image; OCR is helper text only.\n"
            "Never mention OCR, scanning, or handwriting quality.\n"
            "Return JSON only."
        ),
    }

    schema_hint = {
        "page": 2,
        "page_suggestions": ["..."],
        "annotations": [
            {
                "page": 2,
                "type": "language_clarity",
                "rubric_point": "string",
                "anchor_quote": "EXACT substring from OCR_PAGE_TEXT",
                "correction": "string",
                "comment": "string",
            }
        ],
    }

    instructions = (
        "Using the provided annotation rubric text, generate concise actionable annotations for ONE page.\n"
        "Rules:\n"
        "- Prefer 2-5 annotations for the answer page.\n"
        "- Every annotation must be locatable from anchor_quote.\n"
        "- anchor_quote must be an exact contiguous substring from OCR_PAGE_TEXT.\n"
        "- If anchor cannot be found, skip that annotation.\n"
        "- Keep comments concise and corrective.\n"
        "- page_suggestions should be 2-4 concise actionable bullets for this page.\n"
        "- Ignore unrelated watermark/camera/footer artifacts or stray words not part of the answer.\n"
        "- Never mention OCR/scanning/handwriting.\n"
        "Return JSON matching schema."
    )

    os.makedirs("debug_llm", exist_ok=True)
    errors: List[Dict[str, Any]] = []
    annotations: List[Dict[str, Any]] = []
    page_suggestions: List[Dict[str, Any]] = []

    image_by_page = {p.get("page"): p for p in page_images}
    # For precis, prioritize answer page(s): page >= 2.
    ocr_pages = [p for p in (ocr_data.get("pages") or []) if int(p.get("page_number") or 0) >= 2]

    for page in ocr_pages:
        page_num = int(page.get("page_number") or 0)
        if page_num <= 0:
            continue
        ocr_page_text = (page.get("ocr_page_text") or "").strip()
        if not ocr_page_text:
            errors.append({"page": page_num, "error": "Missing ocr_page_text"})
            continue

        payload = {
            "annotations_rubric_text": annotations_rubric_text or "",
            "grading_summary": {
                "total_awarded": grading.get("total_awarded"),
                "criteria": grading.get("criteria", []),
                "reasons_for_low_score": grading.get("reasons_for_low_score", []),
            },
            "ocr_page": {
                "page_number": page_num,
                "ocr_page_text": ocr_page_text,
                "lines": page.get("lines", []),
            },
            "page_image": image_by_page.get(page_num),
            "output_schema": schema_hint,
        }

        parsed: Optional[Dict[str, Any]] = None
        last_err: Optional[Exception] = None
        for _ in range(3):
            try:
                resp = _grok_chat(
                    grok_api_key,
                    messages=[system, {"role": "user", "content": instructions + "\n\nDATA:\n" + json.dumps(payload, ensure_ascii=False)}],
                    model=model,
                    temperature=temperature,
                    max_tokens=2500,
                    timeout=200,
                    max_retries=4,
                )
                content = resp["choices"][0]["message"]["content"]
                parsed = parse_json_with_repair(
                    grok_api_key,
                    content,
                    debug_tag=f"precis_annotations_p{page_num}",
                    max_fix_attempts=2,
                    repair_model=repair_model,
                    repair_temperature=repair_temperature,
                )
                if not isinstance(parsed, dict):
                    raise ValueError("Annotation JSON is not object")
                if not isinstance(parsed.get("annotations"), list):
                    raise ValueError("Missing annotations list")
                if not isinstance(parsed.get("page_suggestions"), list):
                    raise ValueError("Missing page_suggestions list")
                break
            except Exception as e:
                last_err = e
                parsed = None

        if parsed is None:
            errors.append({"page": page_num, "error": str(last_err) if last_err else "unknown"})
            continue

        cleaned_ann: List[Dict[str, Any]] = []
        for a in (parsed.get("annotations") or []):
            if not isinstance(a, dict):
                continue
            aq = str(a.get("anchor_quote", "")).strip()
            if not aq or aq not in ocr_page_text:
                continue
            cleaned_ann.append(
                {
                    "page": page_num,
                    "type": str(a.get("type", "")).strip(),
                    "rubric_point": str(a.get("rubric_point", "")).strip(),
                    "anchor_quote": aq,
                    "target_word_or_sentence": "",
                    "context_before": "",
                    "context_after": "",
                    "correction": str(a.get("correction", "")).strip(),
                    "comment": str(a.get("comment", "")).strip(),
                }
            )
        annotations.extend(cleaned_ann)
        page_suggestions.append({"page": page_num, "suggestions": [str(x).strip() for x in (parsed.get("page_suggestions") or []) if str(x).strip()]})

    return {"annotations": annotations, "page_suggestions": page_suggestions, "errors": errors}

def _dominant_colors_from_scheme(image_path: str) -> Dict[str, Tuple[float, float, float]]:
    fallback = {
        "header_fill": (95 / 255.0, 110 / 255.0, 141 / 255.0),
        "header_text": (1.0, 1.0, 1.0),
        "row_alt": (0.94, 0.95, 0.97),
        "border": (0.35, 0.40, 0.50),
        "section_title": (0.16, 0.20, 0.30),
    }
    if not image_path or not os.path.exists(image_path):
        return fallback
    try:
        img = Image.open(image_path).convert("RGB")
        img.thumbnail((320, 320))
        px = list(img.getdata())

        def _is_whiteish(rgb: Tuple[int, int, int]) -> bool:
            return rgb[0] > 240 and rgb[1] > 240 and rgb[2] > 240

        non_white = [p for p in px if not _is_whiteish(p)]
        if not non_white:
            return fallback

        buckets: Dict[Tuple[int, int, int], int] = {}
        for r, g, b in non_white:
            key = (r // 16 * 16, g // 16 * 16, b // 16 * 16)
            buckets[key] = buckets.get(key, 0) + 1

        sorted_buckets = sorted(buckets.items(), key=lambda kv: kv[1], reverse=True)
        dark = next((k for k, _ in sorted_buckets if (k[0] + k[1] + k[2]) < 360), None)
        mid = next((k for k, _ in sorted_buckets if 360 <= (k[0] + k[1] + k[2]) <= 660), None)

        if dark:
            fallback["header_fill"] = tuple([v / 255.0 for v in dark])  # type: ignore
            fallback["section_title"] = tuple([max(0, min(1, (v - 25) / 255.0)) for v in dark])  # type: ignore
            fallback["border"] = tuple([max(0, min(1, (v - 15) / 255.0)) for v in dark])  # type: ignore
        if mid:
            fallback["row_alt"] = tuple([min(1.0, (v + 45) / 255.0) for v in mid])  # type: ignore
        return fallback
    except Exception:
        return fallback


def _wrap_lines(text: str, fontname: str, fontsize: float, max_width: float) -> List[str]:
    text = (text or "").strip()
    if not text:
        return [""]
    words = text.split()
    lines: List[str] = []
    cur = ""
    for w in words:
        test = (cur + " " + w).strip()
        if fitz.get_text_length(test, fontname=fontname, fontsize=fontsize) <= max_width or not cur:
            cur = test
        else:
            lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    return lines


def _draw_wrapped_text(
    page: fitz.Page,
    x: float,
    y: float,
    text: str,
    *,
    fontname: str,
    fontsize: float,
    max_width: float,
    color: Tuple[float, float, float],
    line_gap: float = 1.3,
) -> float:
    lines = _wrap_lines(text, fontname=fontname, fontsize=fontsize, max_width=max_width)
    for ln in lines:
        page.insert_text((x, y), ln, fontname=fontname, fontsize=fontsize, color=color)
        y += fontsize * line_gap
    return y


def render_precis_report_pdf(
    grading: Dict[str, Any],
    output_pdf_path: str,
    *,
    colouring_scheme_image: str = "",
    question_text: str = "",
    question_statement_text: str = "",
    question_image_path: str = "",
    max_pages: int = 2,
) -> None:
    """Render report on exactly one page by shrinking all text sizes if needed."""
    palette = _dominant_colors_from_scheme(colouring_scheme_image)
    W, H = 595.0, 842.0
    margin = 30.0

    total_awarded = grading.get("total_awarded", 0)
    total_marks = int(round(sum(float(c.get("marks_allocated", 0)) for c in grading.get("criteria", [])))) or 20
    fields: List[Tuple[str, str]] = [
        ("Student Title", str(grading.get("student_title", ""))),
        (
            "Word Counts",
            f"Original: {grading.get('original_passage_word_count', 0)} | Required (1/3): {grading.get('required_precis_word_count', 0)} | Student: {grading.get('student_precis_word_count', 0)}",
        ),
        ("Total Marks", f"{total_awarded}/{total_marks}"),
    ]

    criteria = grading.get("criteria") or []
    reasons = grading.get("reasons_for_low_score") or []
    if not reasons:
        reasons = ["No major weaknesses identified."]
    ideal = grading.get("ideal_precis") or {}
    ideal_title = str(ideal.get("title", "")).strip() or "(Not provided)"
    ideal_text = str(ideal.get("text", "")).strip() or "(Not provided)"

    def _sizes(shrink: int) -> Dict[str, float]:
        base = max(6.0, REPORT_BASE_TEXT_SIZE - float(shrink))
        return {
            "title": base,
            "field_label": base,
            "field_value": base,
            "table_header": base,
            "table_cell": base,
            "section": base,
            "bullet": base,
            "ideal_title": base,
        }

    def _render_once(doc: fitz.Document, shrink: int) -> bool:
        page = doc.new_page(width=W, height=H)
        s = _sizes(shrink)
        y = margin
        usable_h = H - margin

        def _need(h: float) -> bool:
            nonlocal y
            return (y + h) <= usable_h

        def _advance_after_text(font_size: float, extra: float = 0.0) -> None:
            nonlocal y
            y += font_size + extra

        # Title
        if not _need(s["title"] + 8):
            return False
        page.insert_text((margin, y + s["title"]), "Precis Evaluation Report", fontname="hebo", fontsize=s["title"], color=palette["section_title"])
        y += s["title"] + 8

        # Top fields (no bounding boxes)
        label_w = 120
        for label, value in fields:
            max_w = W - 2 * margin - label_w
            lines = _wrap_lines(str(value), "helv", s["field_value"], max_w)
            line_h = s["field_value"] * 1.2
            block_h = max(s["field_label"], line_h * max(1, len(lines))) + 4
            if not _need(block_h):
                return False
            page.insert_text((margin, y + s["field_label"]), f"{label}:", fontname="hebo", fontsize=s["field_label"], color=palette["section_title"])
            yy = y + s["field_value"]
            for ln in lines:
                page.insert_text((margin + label_w, yy), ln, fontname="helv", fontsize=s["field_value"], color=(0, 0, 0))
                yy += line_h
            y += block_h

        y += 4

        # Table
        headers = ["Criterion", "Marks Allocated", "Marks Awarded", "Key Comments"]
        col_w = [205.0, 78.0, 78.0, W - (margin * 2 + 205.0 + 78.0 + 78.0)]
        header_max_lines = max(len(_wrap_lines(h, "hebo", s["table_header"], col_w[i] - 6)) for i, h in enumerate(headers))
        header_h = max(16.0, header_max_lines * (s["table_header"] * 1.05) + 6)
        if not _need(header_h + 2):
            return False
        x = margin
        header_rect = fitz.Rect(margin, y, W - margin, y + header_h)
        page.draw_rect(header_rect, color=palette["border"], fill=palette["header_fill"], width=1)
        for i, htxt in enumerate(headers):
            h_lines = _wrap_lines(htxt, "hebo", s["table_header"], col_w[i] - 6)
            hy = y + s["table_header"] + 1
            for ln in h_lines[:2]:
                page.insert_text((x + 4, hy), ln, fontname="hebo", fontsize=s["table_header"], color=palette["header_text"])
                hy += s["table_header"] * 1.05
            x += col_w[i]
            if i < len(headers) - 1:
                page.draw_line((x, y), (x, y + header_h), color=palette["border"], width=1)
        y += header_h

        for idx, c in enumerate(criteria):
            crit = str(c.get("criterion", ""))
            alloc = str(c.get("marks_allocated", ""))
            award = str(c.get("marks_awarded", ""))
            comment = str(c.get("key_comments", ""))

            crit_lines = _wrap_lines(crit, "helv", s["table_cell"], col_w[0] - 6)
            cmt_lines = _wrap_lines(comment, "helv", s["table_cell"], col_w[3] - 6)
            line_h = s["table_cell"] * 1.2
            row_h = max(18.0, max(len(crit_lines), len(cmt_lines), 1) * line_h + 6)
            if not _need(row_h + 1):
                return False

            fill = palette["row_alt"] if idx % 2 == 0 else (1, 1, 1)
            row_rect = fitz.Rect(margin, y, W - margin, y + row_h)
            page.draw_rect(row_rect, color=palette["border"], fill=fill, width=0.7)
            x = margin

            yy = y + s["table_cell"] + 1
            for ln in crit_lines:
                page.insert_text((x + 4, yy), ln, fontname="helv", fontsize=s["table_cell"], color=(0, 0, 0))
                yy += line_h
            x += col_w[0]
            page.draw_line((x, y), (x, y + row_h), color=palette["border"], width=0.7)
            page.insert_text((x + 4, y + s["table_cell"] + 1), alloc, fontname="helv", fontsize=s["table_cell"], color=(0, 0, 0))
            x += col_w[1]
            page.draw_line((x, y), (x, y + row_h), color=palette["border"], width=0.7)
            page.insert_text((x + 4, y + s["table_cell"] + 1), award, fontname="helv", fontsize=s["table_cell"], color=(0, 0, 0))
            x += col_w[2]
            page.draw_line((x, y), (x, y + row_h), color=palette["border"], width=0.7)
            yy = y + s["table_cell"] + 1
            for ln in cmt_lines:
                page.insert_text((x + 4, yy), ln, fontname="helv", fontsize=s["table_cell"], color=(0, 0, 0))
                yy += line_h
            y += row_h

        y += 6

        # Reasons for low score
        if not _need(s["section"] + 6):
            return False
        page.insert_text((margin, y + s["section"]), "Reasons for Low Score", fontname="hebo", fontsize=s["section"], color=palette["section_title"])
        y += s["section"] + 4

        bullet_lh = s["bullet"] * 1.25
        for item in reasons[:8]:
            blines = _wrap_lines(f"- {str(item)}", "helv", s["bullet"], W - 2 * margin - 10)
            bh = max(1, len(blines)) * bullet_lh + 1
            if not _need(bh):
                return False
            yy = y + s["bullet"]
            for ln in blines:
                page.insert_text((margin + 8, yy), ln, fontname="helv", fontsize=s["bullet"], color=(0, 0, 0))
                yy += bullet_lh
            y += bh

        y += 4

        # Ideal precis (no bounding box)
        if not _need(s["section"] + 6):
            return False
        page.insert_text((margin, y + s["section"]), "Ideal Precis", fontname="hebo", fontsize=s["section"], color=palette["section_title"])
        y += s["section"] + 4

        title_lines = _wrap_lines(f"Title: {ideal_title}", "hebo", s["ideal_title"], W - 2 * margin - 8)
        title_lh = s["ideal_title"] * 1.2
        th = max(1, len(title_lines)) * title_lh + 2
        if not _need(th):
            return False
        yy = y + s["ideal_title"]
        for ln in title_lines:
            page.insert_text((margin + 4, yy), ln, fontname="hebo", fontsize=s["ideal_title"], color=(0, 0, 0))
            yy += title_lh
        y += th

        body_lines = _wrap_lines(ideal_text, "helv", s["bullet"], W - 2 * margin - 8)
        body_lh = s["bullet"] * 1.25
        bh = max(1, len(body_lines)) * body_lh + 2
        if not _need(bh):
            return False
        yy = y + s["bullet"]
        for ln in body_lines:
            page.insert_text((margin + 4, yy), ln, fontname="helv", fontsize=s["bullet"], color=(0, 0, 0))
            yy += body_lh
        return True

    def _render_question_page(doc: fitz.Document) -> None:
        page = doc.new_page(width=W, height=H)
        y = margin
        label_size = REPORT_BASE_TEXT_SIZE
        body_size = max(8.0, REPORT_BASE_TEXT_SIZE - 2)

        if question_text:
            page.insert_text((margin, y + label_size), "Question:", fontname="hebo", fontsize=label_size, color=palette["section_title"])
            y += label_size + 4
            y = _draw_wrapped_text(
                page,
                margin,
                y + body_size,
                question_text,
                fontname="helv",
                fontsize=body_size,
                max_width=W - 2 * margin,
                color=(0, 0, 0),
                line_gap=1.25,
            )
            y += 8

        if question_statement_text:
            page.insert_text((margin, y + label_size), "Passage:", fontname="hebo", fontsize=label_size, color=palette["section_title"])
            y += label_size + 4
            y = _draw_wrapped_text(
                page,
                margin,
                y + body_size,
                question_statement_text,
                fontname="helv",
                fontsize=body_size,
                max_width=W - 2 * margin,
                color=(0, 0, 0),
                line_gap=1.25,
            )
            y += 8

        if question_image_path and os.path.exists(question_image_path):
            try:
                img = Image.open(question_image_path)
                img_w, img_h = img.size
                if img_w > 0 and img_h > 0:
                    avail_w = W - 2 * margin
                    avail_h = H - y - margin
                    scale = min(avail_w / img_w, avail_h / img_h)
                    draw_w = img_w * scale
                    draw_h = img_h * scale
                    x0 = margin + (avail_w - draw_w) / 2.0
                    rect = fitz.Rect(x0, y, x0 + draw_w, y + draw_h)
                    buf = io.BytesIO()
                    img.convert("RGB").save(buf, format="PNG")
                    page.insert_image(rect, stream=buf.getvalue())
            except Exception:
                pass

    # Try progressively smaller global text sizes until it fits one page.
    best_doc: Optional[fitz.Document] = None
    for shrink in range(0, 13):
        d = fitz.open()
        if _render_once(d, shrink):
            best_doc = d
            break
        d.close()

    if best_doc is None:
        # Last fallback at smallest size; keep first page only.
        best_doc = fitz.open()
        _render_once(best_doc, 12)

    if (question_text or question_statement_text or (question_image_path and os.path.exists(question_image_path))):
        _render_question_page(best_doc)

    os.makedirs(os.path.dirname(output_pdf_path) or ".", exist_ok=True)
    best_doc.save(output_pdf_path)
    best_doc.close()

def run_precis_grading(
    pdf_path: str,
    output_json_path: str,
    output_pdf_path: str,
    *,
    rubric_docx: str,
    annotations_rubric_docx: str,
    env_file: str,
    colouring_scheme_image: str,
    extra_json_path: str,
    report_only_pdf_path: Optional[str] = None,
    ocr_workers: int = 2,
    grading_model: str = DEFAULT_MODELS["grading"]["model"],
    grading_temperature: float = float(DEFAULT_MODELS["grading"]["temperature"]),
    annotations_model: str = DEFAULT_MODELS["annotations"]["model"],
    annotations_temperature: float = float(DEFAULT_MODELS["annotations"]["temperature"]),
    repair_model: str = DEFAULT_MODELS["json_repair"]["model"],
    repair_temperature: float = float(DEFAULT_MODELS["json_repair"]["temperature"]),
) -> Dict[str, Any]:
    validate_input_paths(pdf_path, output_json_path, output_pdf_path)
    grok_key, doc_client = load_environment(env_file)
    rubric_text = _load_docx_text(rubric_docx)
    annotations_rubric_text = _load_docx_text(annotations_rubric_docx)
    criteria_template = parse_precis_rubric_criteria(rubric_docx)

    timings: Dict[str, float] = {}
    t0_total = time.perf_counter()

    print("Running OCR on precis PDF...")
    t0 = time.perf_counter()
    ocr_data_raw = run_ocr_on_pdf(doc_client, pdf_path, workers=ocr_workers)
    timings["OCR"] = time.perf_counter() - t0
    print(f"OCR done in {_format_duration(timings['OCR'])}")

    t0 = time.perf_counter()
    ocr_data, extra_text_pack = split_extra_text(ocr_data_raw, pdf_path)
    os.makedirs(os.path.dirname(extra_json_path) or ".", exist_ok=True)
    with open(extra_json_path, "w", encoding="utf-8") as f:
        json.dump(extra_text_pack, f, indent=2, ensure_ascii=False)
    timings["Extra text filtering"] = time.perf_counter() - t0
    print(f"Extra text filtering done in {_format_duration(timings['Extra text filtering'])} "
          f"(removed {extra_text_pack.get('removed_line_count', 0)} lines)")

    print("Preparing page images for Grok...")
    t0 = time.perf_counter()
    page_images = pdf_to_page_images_for_grok(pdf_path, max_pages=2, output_dir="grok_images_precis")
    timings["Image prep"] = time.perf_counter() - t0
    print(f"Image prep done in {_format_duration(timings['Image prep'])}")

    print("Grading precis with rubric...")
    t0 = time.perf_counter()
    grading = call_grok_for_precis_grading(
        grok_key,
        rubric_text=rubric_text,
        criteria_template=criteria_template,
        ocr_data=ocr_data,
        page_images=page_images,
        model=grading_model,
        temperature=grading_temperature,
        repair_model=repair_model,
        repair_temperature=repair_temperature,
    )
    timings["LLM grading"] = time.perf_counter() - t0
    print(f"LLM grading done in {_format_duration(timings['LLM grading'])}")

    print("Generating precis annotations...")
    t0 = time.perf_counter()
    ann_pack = call_grok_for_precis_annotations(
        grok_key,
        annotations_rubric_text=annotations_rubric_text,
        ocr_data=ocr_data,
        grading=grading,
        page_images=page_images,
        model=annotations_model,
        temperature=annotations_temperature,
        repair_model=repair_model,
        repair_temperature=repair_temperature,
    )
    timings["Annotations"] = time.perf_counter() - t0
    print(f"Annotations done in {_format_duration(timings['Annotations'])}")

    annotations = ann_pack.get("annotations") or []
    page_suggestions = ann_pack.get("page_suggestions") or []
    ann_errors = ann_pack.get("errors") or []

    output = {
        "grading": grading,
        "criteria_template": criteria_template,
        "ocr_pages": len(ocr_data.get("pages", [])),
        "annotations": annotations,
        "page_suggestions": page_suggestions,
        "annotation_errors": ann_errors,
        "extra_text_json_path": extra_json_path,
        "model_config": {
            "grading": {"model": grading_model, "temperature": grading_temperature},
            "annotations": {"model": annotations_model, "temperature": annotations_temperature},
            "json_repair": {"model": repair_model, "temperature": repair_temperature},
        },
    }
    os.makedirs(os.path.dirname(output_json_path) or ".", exist_ok=True)
    with open(output_json_path, "w", encoding="utf-8") as f:
        json.dump(output, f, indent=2, ensure_ascii=False)
    print(f"Saved JSON -> {output_json_path}")

    question_block_text, question_statement_block_text = _extract_question_blocks(ocr_data)
    question_json_path = os.path.join(os.path.dirname(output_json_path) or ".", "question_extracted.json")
    with open(question_json_path, "w", encoding="utf-8") as f:
        json.dump(
            {
                "question_text": grading.get("question_text", ""),
                "question_statement_text": grading.get("question_statement_text", ""),
                "question_block_text": question_block_text,
                "question_statement_block_text": question_statement_block_text,
            },
            f,
            indent=2,
            ensure_ascii=False,
        )
    print(f"Saved Question JSON -> {question_json_path}")

    answer_title_hint, answer_block_text = _extract_answer_block_text(ocr_data)
    answer_debug = _LAST_ANSWER_BLOCK_DEBUG if isinstance(_LAST_ANSWER_BLOCK_DEBUG, dict) else {}
    answer_json_path = os.path.join(os.path.dirname(output_json_path) or ".", "answer_extracted.json")
    with open(answer_json_path, "w", encoding="utf-8") as f:
        json.dump(
            {
                "student_title": grading.get("student_title", ""),
                "student_precis_text": grading.get("student_precis_text", ""),
                "answer_title_hint": answer_title_hint,
                "answer_block_text": answer_block_text,
                "fragmented_handwriting": bool(answer_debug.get("fragmented_handwriting", False)),
                "short_ratio": float(answer_debug.get("short_ratio", 0.0) or 0.0),
                "body_y_norm": answer_debug.get("body_y_norm"),
                "line_gap": answer_debug.get("line_gap"),
                "removed_lines": list(answer_debug.get("removed_lines") or []),
                "kept_lines_count": int(answer_debug.get("kept_lines_count", 0) or 0),
                "total_lines_count": int(answer_debug.get("total_lines_count", 0) or 0),
                "used_word_fallback": bool(answer_debug.get("used_word_fallback", False)),
                "word_count_words": int(answer_debug.get("word_count_words", 0) or 0),
                "total_words_count": int(answer_debug.get("total_words_count", 0) or 0),
                "filtered_words_count": int(answer_debug.get("filtered_words_count", 0) or 0),
                "title_source": str(answer_debug.get("title_source", "lines") or "lines"),
            },
            f,
            indent=2,
            ensure_ascii=False,
        )
    print(f"Saved Answer JSON -> {answer_json_path}")

    print("Rendering precis report PDF...")
    t0 = time.perf_counter()
    report_tmp = report_only_pdf_path or os.path.join(os.path.dirname(output_pdf_path) or ".", "_precis_report_tmp.pdf")
    question_text = str(grading.get("question_text", "") or question_block_text or "")
    question_statement_text = str(grading.get("question_statement_text", "") or question_statement_block_text or "")
    question_image_path = ""
    for p in page_images:
        if int(p.get("page") or 0) == 1 and p.get("file_path"):
            question_image_path = str(p.get("file_path"))
            break
    has_question_page = bool(question_text or question_statement_text or question_image_path)
    render_precis_report_pdf(
        grading,
        report_tmp,
        colouring_scheme_image=colouring_scheme_image,
        question_text=question_text,
        question_statement_text=question_statement_text,
        question_image_path=question_image_path,
        max_pages=1,
    )
    timings["PDF render"] = time.perf_counter() - t0
    print(f"PDF render done in {_format_duration(timings['PDF render'])}")

    print("Rendering annotated precis pages...")
    t0 = time.perf_counter()
    if annotate_pdf_essay_pages is None:
        raise RuntimeError("annotate_pdf_with_essay_rubric.py is required for annotation rendering.")

    annotated_pages = annotate_pdf_essay_pages(
        pdf_path=pdf_path,
        ocr_data=ocr_data,
        structure={"outline": {"present": False}, "paragraph_map": []},
        grading=grading,
        annotations=annotations,
        page_suggestions=page_suggestions,
        spelling_errors=None,
        max_callouts_per_page=8,
    )
    if has_question_page and len(annotated_pages) >= 2:
        # Drop the annotated question page to avoid duplicating the report's question page.
        annotated_pages = annotated_pages[1:]
    merge_report_and_annotated_answer(report_tmp, annotated_pages, output_pdf_path)
    if report_only_pdf_path is None:
        try:
            os.unlink(report_tmp)
        except Exception:
            pass
    timings["Merge output PDF"] = time.perf_counter() - t0
    print(f"Merge done in {_format_duration(timings['Merge output PDF'])}")

    total_time = time.perf_counter() - t0_total
    print("\n" + "=" * 60)
    print("PRECIS GRADING TIMING SUMMARY")
    print("=" * 60)
    for k, v in timings.items():
        print(f"  {k}: {_format_duration(v)}")
    print("-" * 60)
    print(f"  Total: {_format_duration(total_time)}")
    print("=" * 60)

    return {
        "status": "success",
        "json_path": output_json_path,
        "pdf_path": output_pdf_path,
        "grading": grading,
        "annotations": annotations,
        "extra_text": extra_text_pack,
        "timings": timings,
        "total_time": total_time,
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="Grade a precis PDF, render report, and append annotated precis pages.")
    parser.add_argument("--pdf", default=os.path.join("precis", "input.pdf"), help="Input precis PDF path")
    parser.add_argument("--output-json", default=os.path.join("precis", "precis_result.json"), help="Output JSON path")
    parser.add_argument("--output-pdf", default=os.path.join("precis", "output.pdf"), help="Output report PDF path")
    parser.add_argument("--rubric-docx", default=os.path.join("precis", "Precis Rubric.docx"), help="Precis rubric .docx path")
    parser.add_argument("--annotations-rubric-docx", default=os.path.join("precis", "ANNOTATIONS RUBRIC FOR PRECIS.docx"), help="Precis annotations rubric .docx path")
    parser.add_argument("--colouring-scheme-image", default=os.path.join("precis", "colouring_scheme.jpeg"), help="Colouring scheme image path")
    parser.add_argument("--env-file", default=os.path.join("precis", "env"), help="Env file path")
    parser.add_argument("--extra-json", default=os.path.join("precis", "extra_text.json"), help="Path to save removed extra/noise OCR text")
    parser.add_argument("--report-only-pdf", default="", help="Optional path to keep standalone report PDF")
    parser.add_argument("--ocr-workers", type=int, default=2, help="Parallel OCR worker count")
    parser.add_argument("--grading-model", default=DEFAULT_MODELS["grading"]["model"])
    parser.add_argument("--grading-temperature", type=float, default=float(DEFAULT_MODELS["grading"]["temperature"]))
    parser.add_argument("--annotations-model", default=DEFAULT_MODELS["annotations"]["model"])
    parser.add_argument("--annotations-temperature", type=float, default=float(DEFAULT_MODELS["annotations"]["temperature"]))
    parser.add_argument("--repair-model", default=DEFAULT_MODELS["json_repair"]["model"])
    parser.add_argument("--repair-temperature", type=float, default=float(DEFAULT_MODELS["json_repair"]["temperature"]))
    args = parser.parse_args()

    result = run_precis_grading(
        pdf_path=args.pdf,
        output_json_path=args.output_json,
        output_pdf_path=args.output_pdf,
        rubric_docx=args.rubric_docx,
        annotations_rubric_docx=args.annotations_rubric_docx,
        env_file=args.env_file,
        colouring_scheme_image=args.colouring_scheme_image,
        extra_json_path=args.extra_json,
        report_only_pdf_path=(args.report_only_pdf or None),
        ocr_workers=args.ocr_workers,
        grading_model=args.grading_model,
        grading_temperature=args.grading_temperature,
        annotations_model=args.annotations_model,
        annotations_temperature=args.annotations_temperature,
        repair_model=args.repair_model,
        repair_temperature=args.repair_temperature,
    )
    print(f"\nDone. Report PDF: {result['pdf_path']}")


if __name__ == "__main__":
    main()
